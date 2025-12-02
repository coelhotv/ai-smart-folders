import os
import shutil
import ollama
import json
import pypdf
import docx
# import pytesseract
# from PIL import Image
from pathlib import Path
import time
import functools
import re
import logging
import subprocess
from typing import Optional, Tuple, Any, List
import shutil as _shutil  # keep name for which checks if needed

# --- 1. AGENT CONFIGURATION ---

# Change these paths to match your setup
# ATTENTION: Use absolute paths to prevent errors

# "Inbox" folder (Where your unorganized files are)
# INBOX_DIR = str(Path.home() / "Dropbox" / "_courses")
INBOX_DIR = str(Path.home() / "testing_ai")

# "Root" folder (Where categorized subfolders will be created)
ORGANIZED_DIR = str(Path.home() / "OrganizedFiles")

# Your locally installed Ollama model (e.g., 'llama3', 'phi3:medium')
OLLAMA_MODEL = 'gemma3:1b'

# Adjust the categories you want the AI to use.
# This helps the AI provide consistent responses.
# EXAMPLE_CATEGORIES = "['Product Management', 'Agile', 'Research', 'Presentations', 'UX', 'Platforms', 'Summaries', 'Programming', 'Misc']"

# --- (V3) Folder names to ignore when scanning for categories ---
IGNORE_DIRS = ['_Unprocessed', '_Ignored'] 

# --- END OF CONFIGURATION ---

# New helper utilities (retry, sanitizers, unique path)
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

def retry(exceptions, tries: int = 3, delay: float = 1.0, backoff: int = 2):
    def decorator(f):
        @functools.wraps(f)
        def wrapper(*args, **kwargs):
            _tries, _delay = tries, delay
            while _tries > 1:
                try:
                    return f(*args, **kwargs)
                except exceptions as e:
                    logger.warning("Transient error: %s. Retrying in %.1fs...", e, _delay)
                    time.sleep(_delay)
                    _tries -= 1
                    _delay *= backoff
            return f(*args, **kwargs)
        return wrapper
    return decorator

def sanitize_category(name: str) -> str:
    """Return a safe folder name from the LLM-provided category."""
    if not name or not isinstance(name, str):
        return "_Unprocessed"
    # remove control chars and reserved file chars, collapse whitespace
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '', name).strip()
    safe = re.sub(r'\s+', ' ', safe)
    # limit length
    return safe[:100] or "_Unprocessed"

def sanitize_filename(name: str, original_ext: str) -> str:
    """Return a safe filename, ensure extension preserved."""
    if not name or not isinstance(name, str):
        base = "file"
    else:
        # remove folder separators & control chars
        base = re.sub(r'[\\/]+', '-', name)
        base = re.sub(r'[\x00-\x1f<>:"|?*]+', '', base).strip()
    # ensure extension
    if not base.lower().endswith(original_ext.lower()):
        base = os.path.splitext(base)[0] + original_ext
    # limit filename length
    return base[:200]

def unique_dest_path(folder: str, filename: str) -> str:
    """Return a non-colliding path inside folder for filename."""
    base, ext = os.path.splitext(filename)
    candidate = filename
    n = 1
    while os.path.exists(os.path.join(folder, candidate)):
        candidate = f"{base}_{n}{ext}"
        n += 1
    return os.path.join(folder, candidate)

def extract_text(file_path: str) -> Optional[str]:
    """
    Extracts raw text from different file types.
    Returns None if the file type is unsupported or an error occurs.
    If a PDF has no extractable text and Tesseract is available, falls back to OCR.
    """
    ext = os.path.splitext(file_path)[1].lower()
    text_parts = []
    try:
        if ext == '.pdf':
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                # page.extract_text() can be None
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
            joined = "\n".join(text_parts).strip()
            if not joined:
                # Try OCR fallback if tesseract exists
                if is_tesseract_available():
                    logger.info("No text found in PDF pages; attempting OCR fallback for %s", file_path)
                    # render each page to image would be heavier; try single-image approach with PIL if possible,
                    # but common robust solution would require pdf2image; skip complex fallback here.
                    try:
                        img_text = pytesseract.image_to_string(Image.open(file_path), lang='eng')
                        return img_text.strip() or None
                    except Exception as e:
                        logger.warning("OCR fallback failed: %s", e)
                        return None
                return None
            return joined

        elif ext == '.docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text_parts.append(para.text)
            return "\n".join(text_parts).strip() or None

        elif ext == '.doc':
            out_dir = os.path.dirname(file_path)
            subprocess.check_call(['soffice', '--headless', '--convert-to', 'docx', '--outdir', out_dir, file_path])

            base = os.path.splitext(file_path)[0]
            docx_path = os.path.join(out_dir, base + '.docx')

            doc = docx.Document(docx_path)
            for para in doc.paragraphs:
                text_parts.append(para.text)
            return "\n".join(text_parts).strip() or None

        # elif ext in ['.jpg', '.jpeg', '.png']:
        #     try:
        #         return pytesseract.image_to_string(Image.open(file_path), lang='eng').strip() or None
        #     except Exception as e:
        #         logger.warning("Image OCR failed for %s: %s", file_path, e)
        #         return None

        elif ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read().strip() or None
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read().strip() or None

        else:
            # Unsupported file type for text extraction
            return None

    except Exception as e:
        logger.warning("Error extracting text from %s: %s", os.path.basename(file_path), e)
        return None

def get_existing_categories(dir_path: str) -> List[str]:
    """
    Scans the target directory and returns a list of existing folder names.
    """
    categories = []
    try:
        for item in os.listdir(dir_path):
            item_path = os.path.join(dir_path, item)
            # Add if it's a directory and not in our ignore list
            if os.path.isdir(item_path) and item not in IGNORE_DIRS:
                categories.append(item)
    except FileNotFoundError:
        # This is normal on the first run
        pass
    except Exception as e:
        print(f"  [!] Error scanning for categories: {e}")
    
    return categories

@retry(Exception, tries=3, delay=1, backoff=2)
def call_ollama_chat(model: str, messages: list, fmt: str = 'json') -> Any:
    """Wrap ollama.chat with retries; bubble exceptions on final failure."""
    # If the local client supports a timeout use it here; otherwise rely on retry/backoff.
    return ollama.chat(model=model, messages=messages, format=fmt)

def get_classification_from_ollama(
    file_content: str,
    filename: str,
    existing_categories: List[str]
) -> Optional[dict]:
    """
    Sends the text and *existing categories* to Ollama for classification.
    Validates and sanitizes the response.
    """
    
    # Advanced Slicing Strategy
    if len(file_content) > 20000:
        # Take the first 15,000 characters and the last 5,000 characters
        content_preview = file_content[:14500] + "\n... [CONTENT SKIPPED] ...\n" + file_content[-5000:]
    else:
        content_preview = file_content

    # content_preview = (file_content[:20000] + '...') if len(file_content) > 20000 else file_content

    prompt = f"""
    You are an expert file organization agent. Your goal is to be consistent.
    A list of existing categories (folders) is: {str(existing_categories)}

    Original Filename: "{filename}"
    Content (sample): "{content_preview}"

    Your task:
    1.  Analyze the file. First, check if it clearly belongs to one of the **existing categories**.
        (If the list is empty, you must propose a new one).
    2.  If it fits an existing category, use that **exact** category name.
    3.  If it does **not** fit *any* existing category, propose a **new**, single, descriptive category name in PascalCase (e.g., 'ProductManagement', 'Research', 'Presentations', 'UX', 'Platforms', 'Summaries', 'Programming', 'Misc').
    4.  Propose a new, descriptive **filename** (use underscores, keep extension).
    
    Respond ONLY with a valid JSON object like this example:
    {{
      "category": "CategoryName",
      "new_filename": "proposed_new_filename.ext"
    }}
    """
    try:
        logger.info("Contacting Ollama (context folders: %d)...", len(existing_categories))
        response = call_ollama_chat(
            model=OLLAMA_MODEL,
            messages=[{'role': 'user', 'content': prompt}],
            fmt='json'
        )

        # Extract content from various response shapes (dicts, objects, ChatResponse)
        content = None
        try:
            if isinstance(response, dict):
                # try both patterns
                if 'message' in response and isinstance(response['message'], dict):
                    content = response['message'].get('content')
                else:
                    content = response.get('content') or response
            else:
                # object-like response (e.g., ollama._types.ChatResponse)
                # try common attributes in order of likelihood
                if hasattr(response, 'message'):
                    msg = getattr(response, 'message')
                    # msg may be dict-like or object with .content
                    if isinstance(msg, dict):
                        content = msg.get('content')
                    elif hasattr(msg, 'content'):
                        content = getattr(msg, 'content')
                if content is None and hasattr(response, 'content'):
                    content = getattr(response, 'content')
                # If content is a JSON string, try parsing it
                if isinstance(content, str):
                    try:
                        parsed = json.loads(content)
                        content = parsed
                    except Exception:
                        # leave as string if not JSON
                        pass
        except Exception as e:
            logger.warning("Failed to normalize Ollama response (%s): %s", type(response), e)
            content = None

        if not isinstance(content, dict):
            logger.warning("Ollama content is not a JSON object/dict.")
            return None

        # Validate fields
        raw_category = content.get('category')
        raw_new_filename = content.get('new_filename')

        if not raw_category or not isinstance(raw_category, str):
            logger.warning("Invalid category returned by Ollama.")
            return None
        if not raw_new_filename or not isinstance(raw_new_filename, str):
            logger.warning("Invalid new_filename returned by Ollama.")
            return None

        # Sanitize outputs
        category = sanitize_category(raw_category)
        original_ext = os.path.splitext(filename)[1] or ''
        new_filename = sanitize_filename(raw_new_filename, original_ext)

        return {"category": category, "new_filename": new_filename}

    except json.JSONDecodeError:
        logger.error("Ollama returned invalid JSON.")
        return None
    except Exception as e:
        logger.error("Ollama API error: %s", e)
        return None

def copy_file(original_path: str, category: str, new_filename: str) -> Optional[str]:
    """
    Copies the original file to the new category folder with the new name.
    Returns destination path on success, None on failure.
    """
    error_dir = os.path.join(ORGANIZED_DIR, "_Unprocessed")

    if not category or not new_filename:
        category = "_Unprocessed"
        new_filename = os.path.basename(original_path)

    # Sanitize category and filename
    safe_category = sanitize_category(category)
    original_ext = os.path.splitext(original_path)[1]
    safe_filename = sanitize_filename(new_filename, original_ext)

    dest_folder = os.path.join(ORGANIZED_DIR, safe_category)
    os.makedirs(dest_folder, exist_ok=True)

    dest_path = unique_dest_path(dest_folder, safe_filename)

    try:
        shutil.copy2(original_path, dest_path)
        logger.info("Copied: %s -> %s/%s", os.path.basename(original_path), safe_category, os.path.basename(dest_path))
        return dest_path
    except Exception as e:
        logger.error("Error copying %s: %s", os.path.basename(original_path), e)
        # try:
        #     os.makedirs(error_dir, exist_ok=True)
        #     shutil.copy2(original_path, os.path.join(error_dir, os.path.basename(original_path)))
        # except Exception:
        #     pass
        return None

def is_tesseract_available() -> bool:
    """Return True if tesseract is installed / available to pytesseract."""
    try:
        # prefer pytesseract API if present
        try:
            _ = pytesseract.get_tesseract_version()
            return True
        except Exception:
            # fallback to checking PATH
            return shutil.which("tesseract") is not None
    except Exception:
        return False

def check_prereqs() -> bool:
    """Ensure Ollama is reachable and Tesseract presence (recommended)."""
    try:
        ollama.list()
    except Exception as e:
        logger.critical("Could not connect to Ollama: %s", e)
        return False

    if not is_tesseract_available():
        logger.warning("Tesseract not found. OCR fallbacks will be disabled. Install tesseract for OCR: brew install tesseract")
    return True

def main():
    """
    Main function to orchestrate the process.
    """
    print(f"--- Starting Organization Agent (V1) ---")
    print(f"Inbox: {INBOX_DIR}")
    print(f"Output: {ORGANIZED_DIR}")
    print(f"AI Model: {OLLAMA_MODEL}\n")

    # Ensure main directories exist
    os.makedirs(INBOX_DIR, exist_ok=True)
    os.makedirs(ORGANIZED_DIR, exist_ok=True)
    # Ensure processed folder exists for originals
    processed_dir = os.path.join(ORGANIZED_DIR, "_Processed")
    os.makedirs(processed_dir, exist_ok=True)

    # --- (NEW V3 STEP) ---
    # Scan for existing categories BEFORE processing files
    print("Scanning for existing categories...")
    existing_categories = get_existing_categories(ORGANIZED_DIR)
    if existing_categories:
        print(f"  [+] Found: {existing_categories}")
    else:
        print("  [*] No existing categories found. Will create new ones.")
    # --- END NEW V3 STEP ---


    files_to_process = [f for f in os.listdir(INBOX_DIR) if os.path.isfile(os.path.join(INBOX_DIR, f)) and not f.startswith('.')]

    if not files_to_process:
        print("No files found to organize.")
        return

    for filename in files_to_process:
        file_path = os.path.join(INBOX_DIR, filename)
        print(f"\n--- Processing: {filename} ---")
        
        # 1. Extract Text
        content = extract_text(file_path)
        
        if not content or content.isspace():
            print("  [!] Empty content or unsupported file type. Leaving it unprocessed.")
            # copy_file(file_path, "_Unprocessed", filename)
            continue
        
        # 2. Classify with AI
        # --- (UPDATED V3 STEP) ---
        # Pass the category list to the classification function
        classification_data = get_classification_from_ollama(
            content, 
            filename, 
            existing_categories
        )
        
        # 3. Move File
        if classification_data and isinstance(classification_data, dict):
            category = classification_data.get('category', 'Misc')
            new_name = classification_data.get('new_filename', filename)
            dest = copy_file(file_path, category, new_name)
            # Move original to _Processed (fallback: append .processed)
            if dest:
                try:
                    shutil.move(file_path, os.path.join(processed_dir, os.path.basename(file_path)))
                except Exception as e:
                    logger.warning("Could not move original to _Processed: %s. Appending .processed", e)
                    try:
                        os.rename(file_path, file_path + ".processed")
                    except Exception:
                        pass
            # --- (NEW V3 STEP) ---
            # Add the new category to our "memory" if it's new
            if category not in existing_categories and category not in IGNORE_DIRS:
                print(f"  [+] New category '{category}' added to context.")
                existing_categories.append(category)
            # --- END NEW V3 STEP ---


        else:
            print("  [!] AI classification failed. Leaving it unprocessed.")
            # copy_file(file_path, "_Unprocessed", filename)

    print("\n--- Organization Complete ---")

if __name__ == "__main__":
    if not check_prereqs():
        print("\n[CRITICAL] Pre-run checks failed. Fix issues then retry.")
    else:
        main()